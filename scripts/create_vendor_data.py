"""
Generate vendor response files: A (xlsx), B (pdf), C (docx), E (txt already exists).
Run once from project root: python scripts/create_vendor_data.py
Requires: pip install fpdf2
"""
import json
import os
import sys

# ── prices ──────────────────────────────────────────────────────────────────
VENDOR_A = [18.50,16.00,24.00,12.50,20.00,42.00,65.00,85.00,38.00,52.00,
             95.00,120.00,145.00,175.00,210.00,28.00,38.00,14.00,18.00,32.00,
             22.00,28.00,35.00,45.00,8.50,55.00,68.00,72.00,88.00,110.00]

VENDOR_B = [17.80,15.40,23.20,12.00,19.20,40.50,62.00,82.00,36.50,50.00,
             91.00,115.00,140.00,168.00,202.00,27.00,36.50,13.50,17.30,30.80,
             21.20,27.00,33.80,43.50,8.20,53.00,65.50,69.50,84.80,106.00]

# Vendor C: lines 1-10 per box, lines 11-20 per 100 units, lines 21-30 per box
VENDOR_C_PROSE = [19.20,16.80,25.00,13.00,20.80,43.50,67.00,88.00,39.50,54.00]
VENDOR_C_PER100 = [9800,12400,15000,18200,22000,2900,3950,1480,1900,3350]
VENDOR_C_BOX2  = [23.00,29.50,36.50,47.00,9.00,57.00,70.00,74.50,91.00,114.00]

with open("data/rfx/RFX-001.json") as f:
    RFX = json.load(f)

ITEMS = RFX["line_items"]

def desc(i):   return ITEMS[i]["description"]
def spec(i):   return ITEMS[i]["spec"]
def qty(i):    return ITEMS[i]["qty"]


# ── Vendor A: clean Excel ────────────────────────────────────────────────────
def make_vendor_a():
    import openpyxl
    from openpyxl.styles import Font, PatternFill, Alignment, Border, Side

    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "RFX-001 Quote Response"

    header_fill = PatternFill("solid", fgColor="1F4E79")
    header_font = Font(color="FFFFFF", bold=True)
    thin = Side(style="thin")
    border = Border(left=thin, right=thin, top=thin, bottom=thin)

    # Company header
    ws.merge_cells("A1:H1")
    ws["A1"] = "VENDOR A PACKAGING SOLUTIONS PVT LTD"
    ws["A1"].font = Font(bold=True, size=14)
    ws.merge_cells("A2:H2")
    ws["A2"] = "Quote Response to RFX-001 | Corrugated Packaging | Dated: 2026-09-05"
    ws.row_dimensions[1].height = 22
    ws.row_dimensions[2].height = 16

    # Column headers (row 4)
    headers = ["Line ID","Description","Spec","Qty","Unit","Unit Price (INR)","GST @18%","Total (INR)"]
    for col, h in enumerate(headers, 1):
        cell = ws.cell(row=4, column=col, value=h)
        cell.fill = header_fill
        cell.font = header_font
        cell.alignment = Alignment(horizontal="center", wrap_text=True)
        cell.border = border
    ws.row_dimensions[4].height = 20

    # Data rows (5..34)
    for i, price in enumerate(VENDOR_A):
        row = 5 + i
        gst = round(price * 0.18, 2)
        total = round(price * qty(i), 2)
        vals = [i+1, desc(i), spec(i), qty(i), "per box", price, gst, total]
        for col, v in enumerate(vals, 1):
            cell = ws.cell(row=row, column=col, value=v)
            cell.border = border
            if col in (6,7,8):
                cell.number_format = '#,##0.00'

    # Questionnaire section
    ws.cell(row=36, column=1, value="QUESTIONNAIRE RESPONSES").font = Font(bold=True, size=11)
    q_data = [
        ("ISO 9001 Certified?", "Yes — Certificate No: ISO9001-A-2024-1234"),
        ("Average Rejection Rate (last 12 months)", "0.5%"),
        ("Lead Time from PO Date", "12 days"),
        ("Manufacturing Location", "Pune, Maharashtra"),
        ("Deviations from Specifications", "None"),
        ("Quote Validity", "30 days"),
    ]
    for j, (q, a) in enumerate(q_data):
        ws.cell(row=37+j, column=1, value=q).font = Font(bold=True)
        ws.cell(row=37+j, column=2, value=a)

    # Column widths
    ws.column_dimensions["A"].width = 10
    ws.column_dimensions["B"].width = 32
    ws.column_dimensions["C"].width = 28
    ws.column_dimensions["D"].width = 8
    ws.column_dimensions["E"].width = 10
    ws.column_dimensions["F"].width = 16
    ws.column_dimensions["G"].width = 12
    ws.column_dimensions["H"].width = 16

    path = "data/vendor_responses/vendor_a_response.xlsx"
    wb.save(path)
    print(f"Created: {path}")


# ── Vendor B: PDF with footnote discount ─────────────────────────────────────
def make_vendor_b():
    try:
        from fpdf import FPDF
    except ImportError:
        print("fpdf2 not found. Run: pip install fpdf2")
        sys.exit(1)

    class PDF(FPDF):
        def header(self):
            self.set_font("Helvetica", "B", 14)
            self.set_text_color(31, 78, 121)
            self.cell(0, 10, "VENDOR B LOGISTICS & PACKAGING LLP", ln=True, align="C")
            self.set_font("Helvetica", "", 9)
            self.set_text_color(80, 80, 80)
            self.cell(0, 6, "Plot 42, MIDC Andheri East, Mumbai 400093  |  gstin: 27AAACV1234F1Z5", ln=True, align="C")
            self.cell(0, 6, "Quote Reference: RFX-001 | Date: 2026-09-06", ln=True, align="C")
            self.ln(4)
            self.set_draw_color(31, 78, 121)
            self.set_line_width(0.5)
            self.line(10, self.get_y(), 200, self.get_y())
            self.ln(3)

        def footer(self):
            self.set_y(-15)
            self.set_font("Helvetica", "I", 8)
            self.set_text_color(128)
            self.cell(0, 10, f"Page {self.page_no()}", align="C")

    pdf = PDF()
    pdf.set_auto_page_break(auto=True, margin=20)
    pdf.add_page()

    # Intro paragraph
    pdf.set_font("Helvetica", "", 10)
    pdf.set_text_color(0)
    pdf.multi_cell(0, 6,
        "Dear Procurement Team,\n\nWe are pleased to submit our competitive quotation for RFX-001 "
        "(Corrugated Packaging). All prices are in Indian Rupees (INR) per box, inclusive of "
        "manufacturing but exclusive of GST. Please refer to the price schedule below.\n\n"
        "We are ISO 9001:2015 certified (Certificate No: ISO9001-B-2023-5678) with a rejection rate "
        "of 1.2% over the last 12 months. We can meet 14-day delivery from PO date from our "
        "Mumbai manufacturing facility.\n")
    pdf.ln(2)

    # Table header
    col_widths = [12, 60, 20, 12, 26, 24, 36]
    col_heads  = ["Line", "Description", "Qty", "Unit", "Unit Price (INR)", "GST @18%", "Total (INR)"]

    def table_header():
        pdf.set_fill_color(31, 78, 121)
        pdf.set_text_color(255, 255, 255)
        pdf.set_font("Helvetica", "B", 8)
        for w, h in zip(col_widths, col_heads):
            pdf.cell(w, 7, h, border=1, align="C", fill=True)
        pdf.ln()
        pdf.set_text_color(0)
        pdf.set_font("Helvetica", "", 8)

    table_header()
    fill = False
    for i, price in enumerate(VENDOR_B):
        gst   = round(price * 0.18, 2)
        total = round(price * qty(i), 2)
        if pdf.get_y() > 255:
            pdf.add_page()
            table_header()
            fill = False
        pdf.set_fill_color(240, 248, 255) if fill else pdf.set_fill_color(255, 255, 255)
        row_vals = [str(i+1), desc(i), str(qty(i)), "per box",
                    f"{price:,.2f}", f"{gst:,.2f}", f"{total:,.2f}"]
        for w, v in zip(col_widths, row_vals):
            pdf.cell(w, 6, v, border=1, fill=True)
        pdf.ln()
        fill = not fill

    # Page 3: Terms + footnote
    pdf.add_page()
    pdf.set_font("Helvetica", "B", 11)
    pdf.cell(0, 8, "Terms and Conditions", ln=True)
    pdf.set_font("Helvetica", "", 10)
    terms = [
        "1. Payment: Net 30 days from invoice date.",
        "2. Delivery: 14 working days from confirmed PO.",
        "3. Validity: This quotation is valid for 30 days.",
        "4. Taxes: GST @18% applicable on all items.",
        "5. Minimum Order: As per quantities specified in RFX.",
        "6. Packaging: Standard export packing.",
        "7. Disputes: Subject to Mumbai jurisdiction.",
    ]
    for t in terms:
        pdf.cell(0, 7, t, ln=True)
    pdf.ln(4)
    pdf.set_font("Helvetica", "B", 10)
    pdf.cell(0, 7, "Quality Assurance", ln=True)
    pdf.set_font("Helvetica", "", 10)
    pdf.multi_cell(0, 7,
        "ISO 9001:2015 Certified (No: ISO9001-B-2023-5678). Average rejection rate 1.2% "
        "(last 12 months). 100% in-house quality inspection before dispatch.")
    pdf.ln(6)

    # Footnote with discount — critical edge case
    pdf.set_y(-45)
    pdf.set_draw_color(180, 180, 180)
    pdf.line(10, pdf.get_y(), 200, pdf.get_y())
    pdf.ln(2)
    pdf.set_font("Helvetica", "I", 8)
    pdf.set_text_color(80, 80, 80)
    pdf.multi_cell(0, 5,
        "* 5% discount applicable on orders above Rs. 2,00,000 (Rupees Two Lakh). "
        "Discount will be adjusted in the final invoice. Prices quoted in the schedule above "
        "are standard list prices and do not reflect this discount.")

    path = "data/vendor_responses/vendor_b_response.pdf"
    pdf.output(path)
    print(f"Created: {path}")


# ── Vendor C: DOCX with prices in prose ──────────────────────────────────────
def make_vendor_c():
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()

    # Cover letter
    h = doc.add_heading("VENDOR C CORRUGATED INDUSTRIES", 0)
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph("Nagpur, Maharashtra | ISO Status: Not Certified")
    doc.add_paragraph("Date: 2026-09-07 | Reference: RFX-001")
    doc.add_paragraph("")

    intro = doc.add_paragraph()
    intro.add_run("Dear Sir/Madam,\n\n").bold = True
    intro.add_run(
        "We thank you for the opportunity to quote for your corrugated packaging requirements "
        "under RFX-001. Please find our rates below. Our average rejection rate over the past "
        "12 months is 2.1%. We can meet the 14-day delivery requirement from our Nagpur facility. "
        "Please note we are not ISO 9001 certified at this time, though our quality management "
        "system meets equivalent standards.\n\n"
        "All prices are in Indian Rupees (INR). No deviations from the specifications listed.\n"
    )

    # Lines 01-10: per box
    doc.add_heading("Section 1: 3-ply and 5-ply Standard Boxes (Items 01–10)", level=2)
    p = doc.add_paragraph()
    for i in range(10):
        p.add_run(
            f"For item {i+1:02d} ({desc(i)}), we quote ₹{VENDOR_C_PROSE[i]:.2f} per box "
            f"for the requested quantity of {qty(i):,} units. "
        )

    # Lines 11-20: per 100 units (THE EDGE CASE)
    doc.add_heading("Section 2: Heavy Duty and Specialty Boxes (Items 11–20)", level=2)
    p = doc.add_paragraph()
    p.add_run(
        "Please note that for the items in this section, our pricing is structured "
        "per 100 units as is our standard practice for bulk orders of this category.\n\n"
    ).italic = True
    per100_items = ITEMS[10:20]
    for i in range(10):
        idx = 10 + i
        raw = VENDOR_C_PER100[i]
        p.add_run(
            f"Item {idx+1:02d} ({desc(idx)}): ₹{raw:,} per 100 units "
            f"(qty required: {qty(idx):,} boxes). "
        )

    # Lines 21-30: per box
    doc.add_heading("Section 3: Inner Packaging and Custom Items (Items 21–30)", level=2)
    p = doc.add_paragraph()
    for i in range(10):
        idx = 20 + i
        p.add_run(
            f"Item {idx+1:02d} ({desc(idx)}): ₹{VENDOR_C_BOX2[i]:.2f} per box, "
            f"qty {qty(idx):,}. "
        )

    doc.add_paragraph("")
    doc.add_paragraph(
        "We look forward to your positive consideration. Kindly revert for any clarification.\n\n"
        "Warm regards,\nVendor C Corrugated Industries"
    )

    path = "data/vendor_responses/vendor_c_response.docx"
    doc.save(path)
    print(f"Created: {path}")


if __name__ == "__main__":
    os.chdir(os.path.dirname(os.path.dirname(os.path.abspath(__file__))))
    make_vendor_a()
    make_vendor_b()
    make_vendor_c()
    print("\nAll vendor files created. Run scripts/generate_vendor_d_image.py for vendor_d.")
